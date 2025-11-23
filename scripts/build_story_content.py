#!/usr/bin/env python3
"""
Utility to convert the branching narrative Word docs into structured JSON
data that the app can consume at runtime.
Also auto-generates puzzle content (Outliers/Grid) based on the narrative text.
"""

from __future__ import annotations

import json
import re
import random
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

try:
    from docx import Document  # type: ignore
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    exit(1)

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT
OUTPUT_PATH = ROOT / "src" / "data" / "storyNarrative.json"

CASE_LETTERS = {1: "A", 2: "B", 3: "C"}

# Words to ignore when generating puzzles
STOPWORDS = {
    "the", "and", "to", "of", "a", "in", "is", "it", "you", "that", "he", "was", "for", "on",
    "are", "with", "as", "i", "his", "they", "be", "at", "one", "have", "this", "from", "or",
    "had", "by", "word", "but", "what", "some", "we", "can", "out", "other", "were", "all",
    "there", "when", "up", "use", "your", "how", "said", "an", "each", "she", "which", "do",
    "their", "time", "if", "will", "way", "about", "many", "then", "them", "would", "write",
    "like", "so", "these", "her", "long", "make", "thing", "see", "him", "two", "has", "look",
    "more", "day", "could", "go", "come", "did", "my", "sound", "no", "most", "number", "who",
    "over", "know", "water", "than", "call", "first", "people", "may", "down", "side", "been",
    "now", "find", "any", "new", "work", "part", "take", "get", "place", "made", "live", "where",
    "after", "back", "little", "only", "round", "man", "year", "came", "show", "every", "good",
    "me", "give", "our", "under", "name", "very", "through", "just", "form", "sentence", "great",
    "think", "say", "help", "low", "line", "differ", "turn", "cause", "much", "mean", "before",
    "move", "right", "boy", "old", "too", "same", "tell", "does", "set", "three", "want", "air",
    "well", "also", "play", "small", "end", "put", "home", "read", "hand", "port", "large",
    "spell", "add", "even", "land", "here", "must", "big", "high", "such", "follow", "act",
    "why", "ask", "men", "change", "went", "light", "kind", "off", "need", "house", "picture",
    "try", "us", "again", "animal", "point", "mother", "world", "near", "build", "self", "earth",
    "father", "head", "stand", "own", "page", "should", "country", "found", "answer", "school",
    "grow", "study", "still", "learn", "plant", "cover", "food", "sun", "four", "between", "state",
    "keep", "eye", "never", "last", "let", "thought", "city", "tree", "cross", "farm", "hard",
    "start", "might", "story", "saw", "far", "sea", "draw", "left", "late", "run", "don't",
    "while", "press", "close", "night", "real", "life", "few", "north", "open", "seem", "together",
    "next", "white", "children", "begin", "got", "walk", "example", "ease", "paper", "group",
    "always", "music", "those", "both", "mark", "often", "letter", "until", "mile", "river",
    "car", "feet", "care", "second", "book", "carry", "took", "science", "eat", "room", "friend",
    "began", "idea", "fish", "mountain", "stop", "once", "base", "hear", "horse", "cut", "sure",
    "watch", "color", "face", "wood", "main", "enough", "plain", "girl", "usual", "young", "ready",
    "above", "ever", "red", "list", "though", "feel", "talk", "bird", "soon", "body", "dog",
    "family", "direct", "pose", "leave", "song", "measure", "door", "product", "black", "short",
    "numeral", "class", "wind", "question", "happen", "complete", "ship", "area", "half", "rock",
    "order", "fire", "south", "problem", "piece", "told", "knew", "pass", "since", "top", "whole",
    "king", "space", "heard", "best", "hour", "better", "true", "during", "hundred", "five",
    "remember", "step", "early", "hold", "west", "ground", "interest", "reach", "fast", "verb",
    "sing", "listen", "six", "table", "travel", "less", "morning", "ten", "simple", "several",
    "vowel", "toward", "war", "lay", "against", "pattern", "slow", "center", "love", "person",
    "money", "serve", "appear", "road", "map", "rain", "rule", "govern", "pull", "cold", "notice",
    "voice", "unit", "power", "town", "fine", "certain", "fly", "fall", "lead", "cry", "dark",
    "machine", "note", "wait", "plan", "figure", "star", "box", "noun", "field", "rest", "correct",
    "able", "pound", "done", "beauty", "drive", "stood", "contain", "front", "teach", "week",
    "final", "gave", "green", "oh", "quick", "develop", "ocean", "warm", "free", "minute", "strong",
    "special", "mind", "behind", "clear", "tail", "produce", "fact", "street", "inch", "multiply",
    "nothing", "course", "stay", "wheel", "full", "force", "blue", "object", "decide", "surface",
    "deep", "moon", "island", "foot", "system", "busy", "test", "record", "boat", "common", "gold",
    "possible", "plane", "stead", "dry", "wonder", "laugh", "thousand", "ago", "ran", "check",
    "game", "shape", "equate", "hot", "miss", "brought", "heat", "snow", "tire", "bring", "yes",
    "distant", "fill", "east", "paint", "language", "among", "grand", "ball", "yet", "wave",
    "drop", "heart", "am", "present", "heavy", "dance", "engine", "position", "arm", "wide",
    "sail", "material", "size", "vary", "settle", "speak", "weight", "general", "ice", "matter",
    "circle", "pair", "include", "divide", "syllable", "felt", "perhaps", "pick", "sudden",
    "count", "square", "reason", "length", "represent", "art", "subject", "region", "energy",
    "hunt", "probable", "bed", "brother", "egg", "ride", "cell", "believe", "fraction", "forest",
    "sit", "race", "window", "store", "summer", "train", "sleep", "prove", "lone", "leg",
    "exercise", "wall", "catch", "mount", "wish", "sky", "board", "joy", "winter", "sat", "written",
    "wild", "instrument", "kept", "glass", "grass", "cow", "job", "edge", "sign", "visit", "past",
    "soft", "fun", "bright", "gas", "weather", "month", "million", "bear", "finish", "happy",
    "hope", "flower", "clothe", "strange", "gone", "jump", "baby", "eight", "village", "meet",
    "root", "buy", "raise", "solve", "metal", "whether", "push", "seven", "paragraph", "third",
    "shall", "held", "hair", "describe", "cook", "floor", "either", "result", "burn", "hill",
    "safe", "cat", "century", "consider", "type", "law", "bit", "coast", "copy", "phrase",
    "silent", "tall", "sand", "soil", "roll", "temperature", "finger", "industry", "value",
    "fight", "lie", "beat", "excite", "natural", "view", "sense", "ear", "else", "quite",
    "broke", "case", "middle", "kill", "son", "lake", "moment", "scale", "loud", "spring",
    "observe", "child", "straight", "consonant", "nation", "dictionary", "milk", "speed",
    "method", "organ", "pay", "age", "section", "dress", "cloud", "surprise", "quiet", "stone",
    "tiny", "climb", "cool", "design", "poor", "lot", "experiment", "bottom", "key", "iron",
    "single", "stick", "flat", "twenty", "skin", "smile", "crease", "hole", "trade", "melody",
    "trip", "office", "receive", "row", "mouth", "exact", "symbol", "die", "least", "trouble",
    "shout", "except", "wrote", "seed", "tone", "join", "suggest", "clean", "break", "lady",
    "yard", "rise", "bad", "blow", "oil", "blood", "touch", "grew", "cent", "mix", "team",
    "wire", "cost", "lost", "brown", "wear", "garden", "equal", "sent", "choose", "fell", "fit",
    "flow", "fair", "bank", "collect", "save", "control", "decimal", "gentle", "woman", "captain",
    "practice", "separate", "difficult", "doctor", "please", "protect", "noon", "whose", "locate",
    "ring", "character", "insect", "caught", "period", "indicate", "radio", "spoke", "atom",
    "human", "history", "effect", "electric", "expect", "crop", "modern", "element", "hit",
    "student", "corner", "party", "supply", "bone", "rail", "imagine", "provide", "agree",
    "thus", "capital", "won't", "chair", "danger", "fruit", "rich", "thick", "soldier", "process",
    "operate", "guess", "necessary", "sharp", "wing", "create", "neighbor", "wash", "bat",
    "rather", "crowd", "corn", "compare", "poem", "string", "bell", "depend", "meat", "rub",
    "tube", "famous", "dollar", "stream", "fear", "sight", "thin", "triangle", "planet", "hurry",
    "chief", "colony", "clock", "mine", "tie", "enter", "major", "fresh", "search", "send",
    "yellow", "gun", "allow", "print", "dead", "spot", "desert", "suit", "current", "lift",
    "rose", "continue", "block", "chart", "hat", "sell", "success", "company", "subtract",
    "event", "particular", "deal", "swim", "term", "opposite", "wife", "shoe", "shoulder",
    "spread", "arrange", "camp", "invent", "cotton", "born", "determine", "quart", "nine",
    "truck", "noise", "level", "chance", "gather", "shop", "stretch", "throw", "shine",
    "property", "column", "molecule", "select", "wrong", "gray", "repeat", "require", "broad",
    "prepare", "salt", "nose", "plural", "anger", "claim", "continent", "oxygen", "sugar",
    "death", "pretty", "skill", "women", "season", "solution", "magnet", "silver", "thank",
    "branch", "match", "suffix", "especially", "fig", "afraid", "huge", "sister", "steel",
    "discuss", "forward", "similar", "guide", "experience", "score", "apple", "bought", "led",
    "pitch", "coat", "mass", "card", "band", "rope", "slip", "win", "dream", "evening",
    "condition", "feed", "tool", "total", "basic", "smell", "valley", "nor", "double", "seat",
    "arrive", "master", "track", "parent", "shore", "division", "sheet", "substance", "favor",
    "connect", "post", "spend", "chord", "fat", "glad", "original", "share", "station", "dad",
    "bread", "charge", "proper", "bar", "offer", "segment", "slave", "duck", "instant", "market",
    "degree", "populate", "chick", "dear", "enemy", "reply", "drink", "occur", "support",
    "speech", "nature", "range", "steam", "motion", "path", "liquid", "log", "meant", "quotient",
    "teeth", "shell", "neck", "bridge"
}

FILLER_WORDS = [
    "FILE", "CASE", "LOCK", "CODE", "TIME", "FACT", "LIES", "TRUE", "DARK", "COLD",
    "RAIN", "CITY", "TOWN", "PORT", "DOCK", "SHIP", "PIER", "YARD", "ROOM", "DOOR",
    "WALL", "HALL", "ROOF", "FLOOR", "KEY", "SAFE", "NOTE", "LIST", "PLAN", "MAPS",
    "GRID", "DATA", "INFO", "TAPE", "FILM", "SHOT", "GLOW", "NEON", "LAMP", "BULB",
    "FUSE", "WIRE", "CORD", "LINE", "PATH", "ROAD", "SIGN", "POST", "MAIL", "SEAL",
    "STAMP", "COIN", "CASH", "BILL", "DEBT", "LOAN", "BOND", "DEAL", "SALE", "SOLD",
    "PAID", "COST", "RATE", "RANK", "ROLE", "TASK", "JOB", "DUTY", "WORK", "HIRE",
    "BOSS", "CHIEF", "HEAD", "LEAD", "CREW", "GANG", "TEAM", "SIDE", "ALLY", "FOE",
    "RIVAL", "HATE", "FEAR", "PAIN", "HURT", "HARM", "KILL", "DEAD", "GONE", "LOST",
    "PAST", "LATE", "SOON", "FAST", "SLOW", "WAIT", "HALT", "STOP", "EXIT", "BACK",
    "AWAY", "NEAR", "HERE", "SEEN", "SAID", "TOLD", "HELD", "KEPT", "TOOK", "GAVE",
    "MADE", "BUILT", "WORN", "TORN", "BROKE", "BENT", "BLUE", "GREY", "RED", "BLACK",
    "WHITE", "GOLD", "IRON", "STEEL", "LEAD", "ZINC", "DUST", "DIRT", "MUD", "SAND",
    "SMOKE", "HAZE", "FOG", "MIST", "RAIN", "SNOW", "WIND", "STORM", "HEAT", "COLD"
]

def extract_keywords(text: str, count: int = 4) -> List[str]:
    if not text:
        return []
    # Normalize text
    text = re.sub(r"[^a-zA-Z\s]", " ", text).lower()
    tokens = text.split()
    
    candidates = []
    seen = set()
    
    for token in tokens:
        if len(token) < 3: continue
        if token in STOPWORDS: continue
        if token in seen: continue
        candidates.append(token.upper())
        seen.add(token)
    
    # Sort by length, descending
    candidates.sort(key=lambda x: len(x), reverse=True)
    return candidates[:count]

def generate_grid(outliers: List[str], count: int = 16) -> List[str]:
    grid = list(outliers)
    while len(grid) < count:
        word = random.choice(FILLER_WORDS)
        if word not in grid:
            grid.append(word)
    random.shuffle(grid)
    return grid

def _normalize_path_token(token: str) -> str:
    token = token.strip()
    token = token.replace("Super-Path", "").replace("super-path", "")
    token = token.replace("Path", "").replace("path", "")
    token = token.strip()
    if not token:
        return token
    # Preserve hyphenated letter codes like A-F-L by stripping hyphens
    if re.fullmatch(r"[A-Za-z0-9\-]+", token):
        token = token.replace("-", "")
    # If there is still whitespace, take the first segment
    token = token.split()[0]
    return token


def sanitize_path_key(token: Optional[str]) -> str:
    if not token:
        return "ROOT"
    primary = _normalize_path_token(token)
    cleaned = re.sub(r"[^A-Za-z0-9]", "", primary).upper()
    return cleaned or "ROOT"


def derive_path_key_from_filename(path: Path) -> str:
    stem = path.stem  # e.g., "Chapter 5 AF"
    parts = stem.split()
    if len(parts) <= 2:
        return "ROOT"
    token = " ".join(parts[2:])
    return sanitize_path_key(token)


@dataclass
class DecisionOption:
    key: str
    title: str
    consequence: Optional[str] = None
    focus: Optional[str] = None
    stats: Optional[str] = None
    outcome: Optional[str] = None
    next_chapter: Optional[int] = None
    next_path_key: Optional[str] = None
    raw_lines: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict:
        return {
            "key": self.key,
            "title": self.title,
            "consequence": self.consequence,
            "focus": self.focus,
            "stats": self.stats,
            "outcome": self.outcome,
            "nextChapter": self.next_chapter,
            "nextPathKey": self.next_path_key,
            "details": self.raw_lines,
        }


@dataclass
class Subchapter:
    chapter: int
    index: int
    path_key: str
    title: str
    bridge_text: Optional[str]
    paragraphs: List[str] = field(default_factory=list)
    decision: Optional[Dict] = None
    previously: Optional[str] = None
    
    # Computed puzzle properties
    board: Optional[Dict] = None

    def case_number(self) -> str:
        letter = CASE_LETTERS.get(self.index)
        if not letter:
            raise ValueError(f"Unexpected subchapter index {self.index} in chapter {self.chapter}")
        return f"{self.chapter:03d}{letter}"

    def to_entry(self) -> Dict:
        entry = {
            "chapter": self.chapter,
            "subchapter": self.index,
            "title": self.title,
            "bridgeText": self.bridge_text,
            "narrative": "\n\n".join(self.paragraphs).strip(),
        }
        if self.decision:
            entry["decision"] = self.decision
        if self.previously:
            entry["previously"] = self.previously
        if self.board:
            entry["board"] = self.board
        return entry


CHAPTER_PATTERN = re.compile(r"Ch(?:apter)?\s*(\d+)", re.IGNORECASE)
PATH_PATTERN = re.compile(r"Path\s+([A-Za-z0-9\-]+)", re.IGNORECASE)
CHAPTER_CODE_PATTERN = re.compile(r"Chapter\s+\d+\s*:\s*([A-Za-z0-9\-]+)", re.IGNORECASE)


def extract_target(line: str) -> (Optional[int], Optional[str]):
    """
    Attempt to parse the target chapter/path key from a line of text.
    """
    chapter = None
    chapter_match = CHAPTER_PATTERN.search(line)
    if chapter_match:
        try:
            chapter = int(chapter_match.group(1))
        except ValueError:
            chapter = None

    path_fragment = None
    path_match = PATH_PATTERN.search(line)
    if path_match:
        path_fragment = path_match.group(1)
    else:
        code_match = CHAPTER_CODE_PATTERN.search(line)
        if code_match:
            path_fragment = code_match.group(1)
        else:
            quoted = re.search(r'"([^"]+)"', line)
            if quoted:
                path_fragment = quoted.group(1)

    path_key = sanitize_path_key(path_fragment)
    if path_fragment is None:
        path_key = "ROOT"
    return chapter, path_key


def extract_previously_block(line: str) -> Optional[str]:
    if "PREVIOUSLY" not in line.upper():
        return None
    match = re.search(r"PREVIOUSLY\s*:?", line, re.IGNORECASE)
    if not match:
        return None
    remainder = line[match.end():].strip()
    if not remainder:
        return None
    if "\n\n" in remainder:
        remainder = remainder.split("\n\n", 1)[0].strip()
    normalized_lines = [
        segment.strip().strip('"\u201c\u201d')
        for segment in remainder.splitlines()
        if segment.strip()
    ]
    normalized = "\n".join(normalized_lines).strip()
    return normalized or None


def parse_docx_file(path: Path) -> List[Subchapter]:
    document = Document(path)
    chapter_number = None
    path_key = derive_path_key_from_filename(path)

    subchapters: List[Subchapter] = []
    current_sub: Optional[Subchapter] = None
    current_bridge: Optional[str] = None
    decision_section: Optional[Dict] = None
    current_option: Optional[DecisionOption] = None
    decision_intro: List[str] = []
    pending_previously: Optional[str] = None

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if not line:
            continue

        previously_fragment = extract_previously_block(line)
        if previously_fragment:
            pending_previously = previously_fragment
            continue

        chapter_match = re.match(r"Chapter\s+(\d+)", line, re.IGNORECASE)
        if chapter_match and chapter_number is None:
            chapter_number = int(chapter_match.group(1))
            continue

        if line.startswith("PUZZLE"):
            # puzzle unlock marker, ignore
            continue

        if line.startswith("Bridge Text:"):
            current_bridge = line.split("Bridge Text:", 1)[1].strip().strip('"')
            continue

        if line.startswith("Subchapter"):
            if current_sub:
                subchapters.append(current_sub)
            sub_header = line
            sub_match = re.match(r"Subchapter\s+(\d+)\.(\d+)", sub_header, re.IGNORECASE)
            if not sub_match:
                raise ValueError(f"Unable to parse subchapter header '{sub_header}' in {path}")
            chap = int(sub_match.group(1))
            index = int(sub_match.group(2))
            title_part = sub_header.split(" - ", 1)[1] if " - " in sub_header else sub_header
            current_sub = Subchapter(
                chapter=chap,
                index=index,
                path_key=path_key,
                title=title_part.strip(),
                bridge_text=current_bridge,
            )
            current_bridge = None
            if pending_previously:
                current_sub.previously = pending_previously
                pending_previously = None
            continue

        # Handle Decision Points (Standard and Variant markers)
        if line == "[DECISION POINT]" or line == "$$DECISION POINT$$":
            if current_sub:
                subchapters.append(current_sub)
                current_sub = None
            decision_section = {"intro": [], "options": []}
            decision_intro = decision_section["intro"]
            current_option = None
            continue

        if (line.startswith("OPTION") or line.startswith("$$OPTION")) and decision_section is not None:
            if current_option:
                decision_section["options"].append(current_option)
            
            clean_line = line.replace("$$OPTION", "OPTION")
            option_match = re.match(r"OPTION\s+([A-Z0-9]+)\s*:\s*(.*)", clean_line)
            if not option_match:
                raise ValueError(f"Malformed option line '{line}' in {path}")
            key = option_match.group(1).strip()
            title = option_match.group(2).strip().strip('"')
            current_option = DecisionOption(key=key, title=title)
            continue

        if decision_section is not None:
            if line.startswith("END CHAPTER") or line.startswith("[PATH"):
                continue
            if current_option is None:
                decision_intro.append(line)
            else:
                if line.startswith("Consequence:"):
                    current_option.consequence = line.split("Consequence:", 1)[1].strip()
                elif line.startswith("Focus:"):
                    current_option.focus = line.split("Focus:", 1)[1].strip()
                elif line.startswith("Stats:"):
                    current_option.stats = line.split("Stats:", 1)[1].strip()
                elif line.startswith("Outcome:"):
                    current_option.outcome = line.split("Outcome:", 1)[1].strip()
                    chapter, target_path = extract_target(line)
                    if chapter and target_path and current_option.next_chapter is None:
                        current_option.next_chapter = chapter
                        current_option.next_path_key = target_path
                else:
                    chapter, target_path = extract_target(line)
                    if chapter and target_path and current_option.next_chapter is None:
                        current_option.next_chapter = chapter
                        current_option.next_path_key = target_path
                    else:
                        current_option.raw_lines.append(line)
            continue

        # Narrative lines
        if current_sub:
            current_sub.paragraphs.append(line)

    if current_option and decision_section:
        decision_section["options"].append(current_option)

    if current_sub:
        subchapters.append(current_sub)

    if decision_section and subchapters:
        subchapters[-1].decision = {
            "intro": decision_section.get("intro", []),
            "options": [option.to_dict() for option in decision_section.get("options", [])],
        }

    if chapter_number is None:
        raise ValueError(f"Failed to determine chapter number for {path}")

    # Sanity-check chapter numbers across subchapters
    for sub in subchapters:
        if sub.chapter != chapter_number:
            sub.chapter = chapter_number
            
        # Auto-Generate Puzzle Content
        # Use bridge text if available, otherwise narrative snippet
        source_text = sub.bridge_text if sub.bridge_text else ("\n".join(sub.paragraphs[:5]) if sub.paragraphs else "")
        outliers = extract_keywords(source_text, 4)
        if not outliers:
            outliers = ["CLUE", "LEAD", "PROOF", "TRUTH"]
        
        # Ensure 4 words
        while len(outliers) < 4:
             w = random.choice(FILLER_WORDS)
             if w not in outliers:
                 outliers.append(w)
        
        grid = generate_grid(outliers)
        
        theme_name = "INVESTIGATION"
        if sub.bridge_text:
            words = sub.bridge_text.split()
            if len(words) > 3:
                theme_name = " ".join(words[:2]).upper()
            else:
                theme_name = sub.bridge_text.upper()
            theme_name = re.sub(r"[^A-Z\s]", "", theme_name).strip()[:15]
            
        sub.board = {
            "outlierWords": outliers,
            "grid": [grid[i:i+4] for i in range(0, 16, 4)],
            "outlierTheme": {
                "name": theme_name,
                "icon": "🔎",
                "summary": (sub.bridge_text[:100] + "...") if sub.bridge_text else "Classified Intel"
            }
        }

    return subchapters


def build_story_dataset() -> Dict:
    case_content: Dict[str, Dict[str, Dict]] = {}

    doc_paths = sorted(
        p for p in DOCS_DIR.glob("*.docx") if "old-story" not in p.as_posix()
    )

    for doc_path in doc_paths:
        try:
            subchapters = parse_docx_file(doc_path)
            for sub in subchapters:
                entry = sub.to_entry()
                case_number = sub.case_number()
                case_bucket = case_content.setdefault(case_number, {})
                key = sub.path_key
                if key in case_bucket:
                    # Warn but overwrite if duplicate path keys in same case (rare)
                    print(f"Warning: Duplicate entry for case {case_number} path {key} ({doc_path}). Overwriting.")
                case_bucket[key] = entry
        except Exception as e:
            print(f"Error processing {doc_path}: {e}")

    return {"caseContent": case_content}


def main():
    print("Building story dataset...")
    dataset = build_story_dataset()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with OUTPUT_PATH.open("w", encoding="utf-8") as fp:
        json.dump(dataset, fp, ensure_ascii=False, indent=2)
    print(f"Wrote story dataset to {OUTPUT_PATH} ({OUTPUT_PATH.stat().st_size} bytes).")


if __name__ == "__main__":
    main()
