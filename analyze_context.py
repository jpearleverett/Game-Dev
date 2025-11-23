import sys
import glob
from docx import Document
import os

files = glob.glob('/workspace/*.docx')
# Pick a few specific ones to look deeper
target_files = [f for f in files if "Chapter 11 APE" in f or "Chapter 2 B" in f]
if not target_files:
    target_files = files[:2]

for f in target_files:
    print(f"--- Analyzing {os.path.basename(f)} ---")
    try:
        doc = Document(f)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if "PUZZLE" in text.upper() or "BRIEFING" in text.upper() or "BRIDGE TEXT" in text.upper():
                print(f"Line {i}: {text}")
                # Print next few lines to see context
                for j in range(1, 4):
                    if i+j < len(doc.paragraphs):
                        next_text = doc.paragraphs[i+j].text.strip()
                        if next_text:
                             print(f"  + {next_text}")
    except Exception as e:
        print(f"Error: {e}")
